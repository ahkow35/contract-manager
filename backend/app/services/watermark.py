"""
Watermark service for HighlightEdit.

Adds a branded footer to every page of DOCX documents for free-tier users.
Uses Word section footers so the stamp appears on every page and survives
DOCX → PDF conversion via Adobe PDF Services.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_BRAND_URL = "https://highlight-edit.nyan.sg"
_HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

# Path to the logo PNG for inline embedding (optional — text-only fallback if absent)
_STATIC_DIR = Path(__file__).parent.parent / "static"
_LOGO_PNG = _STATIC_DIR / "logo.png"


def add_watermark(doc: Document, text: str = "Made with HighlightEdit") -> Document:
    """
    Add a branded footer to every page via Word section footers.

    The footer reads:
        [logo]  Created with HighlightEdit · Free Plan · https://highlight-edit.nyan.sg

    Using section footers ensures the stamp appears on every page in both
    DOCX and PDF (via Adobe PDF Services DOCX→PDF conversion).

    Args:
        doc: python-docx Document object (modified in place)
        text: unused, kept for backwards-compatible signature

    Returns:
        The modified Document object.
    """
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        _write_footer_paragraph(section.footer.paragraphs[0])

    return doc


def _write_footer_paragraph(para) -> None:
    """Write the branded stamp into an existing footer paragraph."""
    # Clear any pre-existing content
    for run in list(para.runs):
        run._r.getparent().remove(run._r)

    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)

    # Optionally embed the logo as a small inline image
    if _LOGO_PNG.exists():
        try:
            run = para.add_run()
            run.add_picture(str(_LOGO_PNG), width=Pt(10))
            para.add_run("  ")
        except Exception:
            pass  # Logo image failed — continue with text-only footer

    run_prefix = para.add_run("Created with ")
    run_prefix.font.size = Pt(9)
    run_prefix.font.color.rgb = RGBColor(107, 114, 128)  # #6B7280

    run_brand = para.add_run("HighlightEdit")
    run_brand.font.size = Pt(9)
    run_brand.font.bold = True
    run_brand.font.color.rgb = RGBColor(55, 65, 81)  # #374151

    run_suffix = para.add_run(" \u00b7 Free Plan \u00b7 ")
    run_suffix.font.size = Pt(9)
    run_suffix.font.color.rgb = RGBColor(107, 114, 128)  # #6B7280

    _add_hyperlink(para, _BRAND_URL, _BRAND_URL)


def _add_hyperlink(para, display_text: str, url: str) -> None:
    """Append a clickable hyperlink run to a paragraph using OOXML."""
    r_id = para.part.relate_to(url, _HYPERLINK_REL, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    # 9pt font (half-points)
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "18")
        rPr.append(el)
    # Gray color matching surrounding text
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "6B7280")
    rPr.append(color)
    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = display_text
    run.append(t)

    hyperlink.append(run)
    para._p.append(hyperlink)
