"""Tests for the footer watermark implementation."""
from docx import Document
from app.services.watermark import add_watermark, _BRAND_URL


def make_simple_doc() -> Document:
    """Create a minimal Document with one paragraph of content."""
    doc = Document()
    doc.add_paragraph("This is the document content.")
    return doc


def test_watermark_appears_in_section_footer():
    """The watermark should appear in the section footer (every page), not the body."""
    doc = make_simple_doc()
    add_watermark(doc)
    for section in doc.sections:
        # Collect text from regular runs
        run_text = "".join(
            run.text for para in section.footer.paragraphs for run in para.runs
        )
        # Collect text from hyperlink elements (w:hyperlink > w:r > w:t)
        footer_xml = section.footer._element.xml
        assert "HighlightEdit" in run_text
        assert "Free Plan" in run_text
        assert _BRAND_URL in footer_xml, "Hyperlink URL not found in footer XML"


def test_watermark_does_not_use_header():
    """The document header should NOT contain VML watermark shapes."""
    doc = make_simple_doc()
    add_watermark(doc)
    for section in doc.sections:
        header_xml = section.header._element.xml
        assert "PowerPlusWaterMarkObject" not in header_xml, \
            "VML watermark shape found in header — old watermark not removed"


def test_original_content_preserved():
    """The original document content should still be present after watermarking."""
    doc = make_simple_doc()
    add_watermark(doc)
    texts = [p.text for p in doc.paragraphs]
    assert any("document content" in t for t in texts)
